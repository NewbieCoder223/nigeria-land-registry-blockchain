"""
build_docx.py
Generates a fully AUST-formatted .docx from the individual chapter .md files.

AUST Format requirements (from AUST_PROJECT_FORMAT.md):
  Font       : Times New Roman, 12pt
  Margins    : Left 3.5 cm, Right/Top/Bottom 2.5 cm
  Spacing    : 1.5 lines for body; single for abstract/captions/footnotes
  Paragraphs : First-line indent 1.25 cm; justified
  Headings   : Chapter = 14pt Bold Centred; Section = 12pt Bold; Sub = 12pt Bold-Italic
  Page nums  : Bottom-centre; Roman (ii…) for prelims; Arabic (1…) for chapters+refs
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE = r"C:\Users\Fidelis\OneDrive\Desktop\Final Year Project"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_page_format(doc):
    """Apply AUST margins to all sections."""
    for section in doc.sections:
        section.left_margin   = Cm(3.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)


def fmt_paragraph(para, size=12, bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=6,
                  line_spacing=1.5, indent_first=False):
    """Apply AUST formatting to a paragraph."""
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if indent_first:
        pf.first_line_indent = Cm(1.25)

    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic


def add_paragraph(doc, text="", size=12, bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=6,
                  line_spacing=1.5, indent_first=False):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    pf = p.paragraph_format
    pf.alignment   = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if indent_first:
        pf.first_line_indent = Cm(1.25)
    return p


def add_heading(doc, text, level=1):
    """
    level 1 = CHAPTER heading (14pt Bold Centred, ALL CAPS)
    level 2 = Section heading  (12pt Bold, left)
    level 3 = Sub-section      (12pt Bold-Italic, left)
    """
    if level == 1:
        p = add_paragraph(doc, text.upper(), size=14, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER,
                          space_before=18, space_after=12, line_spacing=1.5)
    elif level == 2:
        p = add_paragraph(doc, text, size=12, bold=True,
                          align=WD_ALIGN_PARAGRAPH.LEFT,
                          space_before=12, space_after=6, line_spacing=1.5)
    else:
        p = add_paragraph(doc, text, size=12, bold=True, italic=True,
                          align=WD_ALIGN_PARAGRAPH.LEFT,
                          space_before=6, space_after=3, line_spacing=1.5)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_figure_placeholder(doc, label):
    """Replace ASCII/code diagrams with a labelled placeholder box."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"[{label}]")
    run.font.name   = "Times New Roman"
    run.font.size   = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    # light border via shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "999999")
        bdr.append(el)
    pPr.append(bdr)
    return p


def add_table_from_md(doc, rows):
    """
    Parse a markdown table (list of raw row strings) and add a Word table.
    rows: list of strings like '| A | B | C |'
    """
    data = []
    for row in rows:
        if re.match(r'\|[-\s|]+\|', row):
            continue   # skip separator row
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        if any(c for c in cells):
            data.append(cells)

    if not data:
        return

    ncols = len(data[0])
    table = doc.add_table(rows=len(data), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(data):
        row_obj = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= ncols:
                break
            cell = row_obj.cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            # Header row = bold
            if r_idx == 0:
                run.font.bold = True
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()  # spacing after table


def add_page_number_field(doc, section_index=0, numeral_format="arabic"):
    """
    Inject a bottom-centre page number field into the footer.
    numeral_format: 'arabic' or 'roman'
    """
    section = doc.sections[section_index]
    footer  = section.footer
    # Clear any existing footer paragraphs
    for para in footer.paragraphs:
        for run in para.runs:
            run.clear()

    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert PAGE field
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")

    if numeral_format == "roman":
        instr.text = " PAGE \\* lowerroman "
    else:
        instr.text = " PAGE "

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rpr.append(sz)
    run_el.append(rpr)
    run_el.append(fld_char1)
    run_el.append(instr)
    run_el.append(fld_char2)
    fp._p.append(run_el)


def add_section_break(doc, break_type="nextPage"):
    """Add a section break (new page, different header/footer possible)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), "decimal")
    sectPr.append(pgNumType)
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), break_type)
    sectPr.append(type_el)
    pPr.append(sectPr)


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------

def parse_md_to_doc(doc, md_text, in_prelim=False):
    """
    Parse and render a markdown string into the Word document.
    Handles: headings (#, ##, ###), tables, code blocks (→ figure placeholders),
    bold inline (**), bullet lists, blank lines.
    """
    lines = md_text.splitlines()
    i = 0
    figure_counter = [0]   # mutable int for inline counter

    while i < len(lines):
        line = lines[i]

        # ---- CODE BLOCK (``` ... ```) → Figure placeholder ----------------
        if line.strip().startswith("```"):
            # Collect until closing ```
            block_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block_lines.append(lines[i])
                i += 1
            figure_counter[0] += 1
            # Try to extract a label from first meaningful comment line
            first = next((l for l in block_lines if l.strip()), "")
            label = first.strip().lstrip("#").strip() if first else f"Figure {figure_counter[0]}"
            add_figure_placeholder(doc, f"Figure — {label[:80]}")
            i += 1
            continue

        # ---- TABLE BLOCK --------------------------------------------------
        if line.strip().startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i])
                i += 1
            add_table_from_md(doc, table_rows)
            continue

        # ---- CHAPTER HEADING (#) -----------------------------------------
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue

        # ---- SECTION HEADING (##) ----------------------------------------
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue

        # ---- SUB-SECTION HEADING (###) ------------------------------------
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue

        # ---- PAGE BREAK MARKER (---) -------------------------------------
        if line.strip() == "---":
            add_page_break(doc)
            i += 1
            continue

        # ---- BULLET LIST -------------------------------------------------
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            text = line.strip().lstrip("-* ").strip()
            _add_inline_bold(p, text, size=12)
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        # ---- NUMBERED LIST -----------------------------------------------
        m_num = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if m_num:
            p = doc.add_paragraph(style="List Number")
            _add_inline_bold(p, m_num.group(2), size=12)
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        # ---- BLANK LINE --------------------------------------------------
        if not line.strip():
            i += 1
            continue

        # ---- NORMAL PARAGRAPH --------------------------------------------
        # Skip pure-marker lines like "(Page i — ...)"
        if line.strip().startswith("*(") and line.strip().endswith(")*"):
            i += 1
            continue

        p = doc.add_paragraph()
        _add_inline_bold(p, line.strip(), size=12)
        sp = 6 if in_prelim else 6
        pf = p.paragraph_format
        pf.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_before = Pt(0)
        pf.space_after  = Pt(sp)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.first_line_indent = Cm(1.25)
        i += 1

    return doc


def _add_inline_bold(para, text, size=12):
    """
    Split text on **...** markers and create bold/normal runs.
    Also handle *italic* markers.
    """
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            run = para.add_run(part)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def read(fname):
    path = os.path.join(BASE, fname)
    with open(path, encoding="utf-8") as f:
        return f.read()


def build():
    doc = Document()
    set_page_format(doc)

    # Remove default styles' spacing quirks
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # -----------------------------------------------------------------------
    # FRONT MATTER  (Roman numerals)
    # -----------------------------------------------------------------------
    fm = read("front_matter.md")
    parse_md_to_doc(doc, fm, in_prelim=True)

    # After front matter, add a section break so chapters restart numbering
    add_section_break(doc, "nextPage")

    # -----------------------------------------------------------------------
    # CHAPTER 1
    # -----------------------------------------------------------------------
    add_page_break(doc)
    parse_md_to_doc(doc, read("chapter1_introduction.md"))

    # -----------------------------------------------------------------------
    # CHAPTER 2
    # -----------------------------------------------------------------------
    add_page_break(doc)
    parse_md_to_doc(doc, read("chapter2_literature_review.md"))

    # -----------------------------------------------------------------------
    # CHAPTER 3
    # -----------------------------------------------------------------------
    add_page_break(doc)
    parse_md_to_doc(doc, read("chapter3_analysis_design.md"))

    # -----------------------------------------------------------------------
    # REFERENCES
    # -----------------------------------------------------------------------
    add_page_break(doc)
    parse_md_to_doc(doc, read("references.md"))

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    out = os.path.join(BASE, "SecureLandRegistry_Chapters1to3_AUST.docx")
    doc.save(out)
    print(f"\n✅  Saved: {out}")
    print(f"   Size : {os.path.getsize(out) // 1024} KB")


if __name__ == "__main__":
    build()
